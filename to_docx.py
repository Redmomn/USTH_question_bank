from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
import json
import os
os.chdir(os.path.dirname(os.path.abspath(__file__)))

filePath = './1003毛概'
jsonFileList = ["单选.json","多选.json","判断.json"]

# 保存docs的文件夹名称
wordPath = ''

docxPath1 = f'./word/{wordPath}'
docsPath2 = f'/docs/.vuepress/public/word/{wordPath}'
os.makedirs(docxPath1)

# Document('./word/test1.docx')

# document = Document()

# paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
# prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
# paragraph.insert_paragraph_before('Lorem ipsum')
# paragraph.insert_paragraph_before('Lorem ipsum')


# document.save(docxPath)

# 原始json整理，返回题目列表和题目类型，单选为1，多选为2，判断为3
def rawJsonSwitch(file):
    raw = json.load(fp=file)
    json_2 = raw['text']
    json_2 = json.loads(json_2)
    if 'BC_xcqm_tkb_SC' in json_2['BusinessData'].keys():
        return json_2['BusinessData']['BC_xcqm_tkb_SC'],1
    if 'BC_xcqm_tkb_MC' in json_2['BusinessData'].keys():
        return json_2['BusinessData']['BC_xcqm_tkb_MC'],2
    if 'BC_xcqm_tkb_TF' in json_2['BusinessData'].keys():
        return json_2['BusinessData']['BC_xcqm_tkb_TF'],3

# 解析题目，写入文件对象
def titWrite(data:list,docxPath):
    document = Document()
    for item in data:
    # 提取题目信息
        xh = item["xh"]
        wt = item["wt"]
        da = item["da"]
        options = [item["xx1"], item["xx2"], item["xx3"], item["xx4"], item["xx5"]]
        
        # 添加题目序号和题目内容
        document.add_paragraph(f'{xh}. {wt}')
        
        if options == ["", "", "", "", ""] and da in ["对", "错"]:
            # 判断题
            if da == "对":
                document.add_paragraph("    对").runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 对答案标红
                document.add_paragraph("    错")
            else:
                document.add_paragraph("    对")
                document.add_paragraph("    错").runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 错答案标红
        else:
            # 其他题型（多选题、单选题等）
            if len(da) == 1:
                # 单选题
                for i, option in enumerate(options):
                    if option:
                        if da == option[0]:
                            # 使用红色标记答案
                            document.add_paragraph(f'   {chr(65 + i)}. {option[2:]}').runs[0].font.color.rgb = RGBColor(255, 0, 0)
                        else:
                            document.add_paragraph(f'   {chr(65 + i)}. {option[2:]}')
            else:
                # 多选题
                for i, option in enumerate(options):
                    if option:
                        if chr(65 + i) in da:
                            # 使用红色标记答案
                            document.add_paragraph(f'   {chr(65 + i)}. {option[2:]}').runs[0].font.color.rgb = RGBColor(255, 0, 0)
                        else:
                            document.add_paragraph(f'   {chr(65 + i)}. {option[2:]}')
    
    # 改变文档为宋体
    document.styles['Normal'].font.name = u'宋体'
    # document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    document.save(docxPath)

def main():
    for jsonFile in jsonFileList:
        with open(file=f'{filePath}/{jsonFile}',mode='r',encoding='utf-8') as file:
            titleList,titleType = rawJsonSwitch(file)
        if titleType == 1:
            docxPath = f'{docxPath1}/单选题.docx'
        elif titleType == 2:
            docxPath = f'{docxPath1}/多选题.docx'
        elif titleType == 3:
            docxPath = f'{docxPath1}/判断题.docx'
        titWrite(titleList,docxPath)

if __name__ == '__main__':
    main()
